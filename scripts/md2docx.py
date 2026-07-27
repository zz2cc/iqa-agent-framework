# -*- coding: utf-8 -*-
"""轻量 Markdown → Word(.docx) 转换器（本项目文档专用）

支持：#-#### 标题、表格、围栏代码块、有序/无序列表、引用、**加粗**、`行内代码`、--- 分隔线。
用法： python scripts/md2docx.py <input.md> [output.docx]
"""
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

EAST_FONT = "微软雅黑"
CODE_FONT = "Consolas"


def set_east_asia(style_or_run, font=EAST_FONT):
    rpr = style_or_run.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def add_runs(par, text, base_bold=False):
    """解析 **bold** 与 `code` 行内标记。"""
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = CODE_FONT
            r.font.size = Pt(9)
        else:
            r = par.add_run(tok)
            r.bold = base_bold
        set_east_asia(r)


def flush_table(doc, rows):
    cells = [[c.strip() for c in re.split(r"(?<!\\)\|", r)[1:-1]] for r in rows]
    header, body = cells[0], [r for r in cells[1:] if not set("".join(r)) <= set("-: ")]
    t = doc.add_table(rows=len(body) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        p = t.rows[0].cells[j].paragraphs[0]
        add_runs(p, h, base_bold=True)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            p = t.rows[i].cells[j].paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "")
    doc.add_paragraph()


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    set_east_asia(normal)
    for i in range(1, 5):
        st = doc.styles[f"Heading {i}"]
        st.font.color.rgb = RGBColor(0x1F, 0x3F, 0x66)
        set_east_asia(st)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()

        if s.startswith("```"):  # 代码块
            buf, i = [], i + 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            for cl in buf:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Pt(12)
                r = p.add_run(cl if cl else " ")
                r.font.name = CODE_FONT
                r.font.size = Pt(9)
            doc.add_paragraph()
            continue

        if s.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
            rows = [s]
            i += 1
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            flush_table(doc, rows)
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", s)
        if m:
            h = doc.add_heading(level=len(m.group(1)))
            add_runs(h, m.group(2))
        elif s == "---":
            pass
        elif s.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, s.lstrip("> "))
        elif re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", s))
        elif re.match(r"^\d+[.、]\s*", s):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+[.、]\s*", "", s))
        elif s:
            p = doc.add_paragraph()
            add_runs(p, s)
        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"
    out = convert(src, dst)
    d = Document(out)  # 回读校验
    print(f"OK: {out} | {os.path.getsize(out)//1024}KB | paragraphs={len(d.paragraphs)} tables={len(d.tables)}")
